import csv
from docx import Document
import re

def remove_links_from_markdown(md: str) -> str:
    if not md:
        return md

    # Remove inline markdown links but keep text
    md = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', md)

    # Remove reference-style links like: [text][1]
    md = re.sub(r'\[([^\]]+)\]\[[^\]]+\]', r'\1', md)

    # Remove bare URLs
    md = re.sub(r'https?://\S+', '', md)

    return md.strip()


def write_to_csv(data, filename="data"):
    filename = f'{filename}.csv'
    if not data:
        print("No data to write")
        return

    # Collect all possible keys across objects
    keys = set()
    for row in data:
        keys.update(row.keys())
    keys = list(keys)

    with open(filename, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=keys)
        writer.writeheader()
        writer.writerows(data)

    print(f"Saved {len(data)} rows to {filename}")




def write_to_word(data, filename="data"):
    filename = f'{filename}.docx'
    if not data:
        print("No data to write")
        return

    doc = Document()
    doc.add_heading("Crawl4AI Webpage Data", level=1)

    for idx, row in enumerate(data, start=1):
        doc.add_heading(f"Page {idx}", level=2)

        for key, value in row.items():
            doc.add_paragraph(f"{key.upper()}:", style="List Bullet")
            doc.add_paragraph(str(value))

        doc.add_page_break()

    doc.save(filename)
    print(f"Saved {len(data)} pages to {filename}")


def write_to_word_clean(data, filename="data"):
    filename = f'{filename}.docx'
    doc = Document()

    for row in data:
        doc.add_heading(row.get("title","Untitled"), level=2)
        doc.add_paragraph(row.get("markdown",""))
        doc.add_page_break()

    doc.save(filename)
